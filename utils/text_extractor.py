"""
Text extraction module for different file formats
"""

import io
import re
import unicodedata
from typing import Union
from docx import Document
from PyPDF2 import PdfReader


def extract_text_from_file(file) -> str:
    """
    Extract text content from uploaded file with normalization
    Applies UTF-8 encoding, NFKC normalization, and cleanup per Final Plan
    
    Args:
        file: Uploaded file object from Streamlit
        
    Returns:
        str: Extracted and normalized text content
    """
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'txt':
        raw_text = extract_from_txt(file)
    elif file_extension == 'docx':
        raw_text = extract_from_docx(file)
    elif file_extension == 'pdf':
        raw_text = extract_from_pdf(file)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    # Apply normalization per Final Plan requirements
    return normalize_text(raw_text)


def extract_from_txt(file) -> str:
    """Extract text from .txt file"""
    try:
        # Try UTF-8 first
        content = file.read().decode('utf-8')
    except UnicodeDecodeError:
        # Fallback to latin-1
        file.seek(0)
        content = file.read().decode('latin-1')
    
    return content


def extract_from_docx(file) -> str:
    """Extract text from .docx file"""
    doc = Document(file)
    
    # Extract text from paragraphs
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    
    return '\n'.join(text_parts)


def extract_from_pdf(file) -> str:
    """Extract text from .pdf file"""
    pdf_reader = PdfReader(file)
    
    text_parts = []
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text.strip():
            text_parts.append(text)
    
    return '\n'.join(text_parts)


def normalize_text(text: str) -> str:
    """
    Normalize extracted text per Final Plan requirements:
    - Ensure UTF-8 encoding
    - Apply NFKC Unicode normalization
    - Remove non-printable control characters
    - Normalize whitespace
    - Preserve paragraph breaks
    
    Args:
        text: Raw text content
        
    Returns:
        str: Normalized text
    """
    # Apply NFKC normalization (compatibility decomposition + canonical composition)
    text = unicodedata.normalize('NFKC', text)
    
    # Remove non-printable control characters (except newlines, tabs, carriage returns)
    text = ''.join(char for char in text if char.isprintable() or char in '\n\t\r')
    
    # Normalize newlines to \n
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Collapse excessive whitespace within lines (preserve paragraph breaks)
    lines = text.split('\n')
    lines = [re.sub(r'[ \t]+', ' ', line.strip()) for line in lines]
    
    # Preserve paragraph breaks (double newline)
    text = '\n'.join(lines)
    
    # Collapse more than 2 consecutive newlines to 2 (preserve paragraph structure)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def clean_text(text: str) -> str:
    """
    Legacy function - now just calls normalize_text
    Kept for backwards compatibility
    
    Args:
        text: Raw text content
        
    Returns:
        str: Cleaned text
    """
    return normalize_text(text)
